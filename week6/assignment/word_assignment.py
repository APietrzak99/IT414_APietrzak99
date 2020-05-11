import docx
import openpyxl
from bs4 import BeautifulSoup
from openpyxl.utils import *
import PyPDF4
import re

text_block = docx.Document("text_files/operations.docx")
all_paragraphs = text_block.paragraphs


my_document = docx.Document()
my_document.add_heading("Operations \n", 1)
for paragraph in all_paragraphs:
    my_document.add_paragraph(paragraph.text)


my_document.add_page_break()

my_document.add_heading("Sales \n", 1)

sales_workbook = openpyxl.load_workbook("text_files/sales.xlsx")

curr_sheet = sales_workbook["Sheet1"]

max_row = curr_sheet.max_row
max_col = curr_sheet.max_column

last_col_letter = get_column_letter(max_col)

data_list = []

for row in curr_sheet.iter_rows(values_only=True):
    data_list.append({"data0":str(row[0]), "data1":str(row[1]),"data2":str(row[2]),"data3":str(row[3]),"data4":str(row[4]),"data5":str(row[5])})
        
data_list_table = my_document.add_table(rows=len(data_list), cols=len(data_list[0]))

record_count = 0

while record_count < len(data_list):
    tmp_row_cells = data_list_table.rows[record_count].cells
    tmp_row_cells[0].text = data_list[record_count]["data0"]
    tmp_row_cells[1].text = data_list[record_count]["data1"]
    tmp_row_cells[2].text = data_list[record_count]["data2"]
    tmp_row_cells[3].text = data_list[record_count]["data3"]
    tmp_row_cells[4].text = data_list[record_count]["data4"]
    tmp_row_cells[5].text = data_list[record_count]["data5"]
    record_count = record_count + 1
    
my_document.add_heading("Marketing \n", 1)

pdfFileObj = open("text_files/marketing.pdf", "rb")
pdfReader = PyPDF4.PdfFileReader(pdfFileObj)
pageObj = pdfReader.getPage(0)

pdfText = pageObj.extractText()

new_text = pdfText.replace('\r', '\n').replace('\n', '')

my_document.add_paragraph(str(new_text))

my_document.add_page_break()

my_document.add_heading("IT \n", 1)

html = open("text_files/IT.html", "rb")
make_soup = BeautifulSoup(html, "lxml")
text = make_soup.get_text()
text= text.strip()
new_text = text.split("\n")

for item in new_text:
    my_document.add_paragraph(item)

my_document.save("text_files/ceo_report.docx")