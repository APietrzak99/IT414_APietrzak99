import docx

text_block = """It was hard to know who missed Nancy more-Tom or the children. He sat around the cabin looking cross and glum. The ground was frozen, so very little work could be done on the farm. He decided, when Andrew Crawford started his school that Abe and Sally might as well go. There was nothing else for them to do, and Nancy would have wanted it.

For the first time since his mother's death Abe seemed to cheer up. Every morning, except when there were chores to do at home, he and Sally took a path through the woods to the log schoolhouse. Master Crawford kept a "blab" school. The "scholars," as he called his pupils, studied their lessons out loud. The louder they shouted, the better he liked it. If a scholar didn't know his lesson, he had to stand in the corner with a long pointed cap on his head. This was called a dunce cap.

One boy who never had to wear a dunce cap was Abe Lincoln. He was too smart. His side won nearly every spelling match. He was good at figuring, and he had the best handwriting of anyone at school. Master Crawford taught reading from the Bible, but he had several other books from which he read aloud. Among Abe's favorite stories were the ones about some wise animals that talked. They were by a man named Aesop who had lived hundreds of years before.

Abe even made up compositions of his own. He called them "sentences." One day he found some of the boys being cruel to a terrapin, or turtle. He made them stop. Then he wrote a composition in which he said that animals had feelings the same as folks."""

text_pieces = text_block.split("\n\n")

para_count = 0

my_document = docx.Document()

while para_count < len(text_pieces):
    my_document.add_heading("Paragraph " + str(para_count + 1), 1)
    tmp_paragraph = my_document.add_paragraph()
    split_text = "Abe"
    run_pieces = text_pieces[para_count].split(split_text)

    run_count = 0
    while run_count < len(run_pieces):
        tmp_paragraph.add_run(run_pieces[run_count])

        if run_count != len(run_pieces) - 1:

            if run_pieces[run_count + 1][0:2] == "'s":
                split_text = "Abe's"
                run_pieces[run_count + 1] = run_pieces[run_count + 1][2:]

            if run_pieces[run_count + 1][0:8] == " Lincoln":
                split_text = "Abe Lincoln"
                run_pieces[run_count + 1] = run_pieces[run_count + 1][8:]
                
            tmp_text = tmp_paragraph.add_run(split_text)
            tmp_text.bold = True

        run_count = run_count + 1

    para_count = para_count + 1

my_document.add_page_break()

reading_list = [{"author": "Charles Dickens", "book": "Oliver Twist"},
                {"author": "Mark Twain", "book": "The Adventures of Tom Sawyer"},
                {"author": "Jules Verne", "book": "Journey to the Center of the Earth"}]

reading_list_table = my_document.add_table(rows=len(reading_list), cols=len(reading_list[0]))

record_count = 0

while record_count < len(reading_list):
    tmp_row_cells = reading_list_table.rows[record_count].cells
    tmp_row_cells[0].text = reading_list[record_count]["author"]
    tmp_row_cells[1].text = reading_list[record_count]["book"]
    record_count = record_count + 1


my_document.save("my_output.docx")